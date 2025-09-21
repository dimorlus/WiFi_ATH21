"""
DOCX Generator Module
Generates DOCX files based on template with sensor data
"""

import os
import shutil
import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from date_formatter import DateFormatter


class DOCXGenerator:
    """Generate DOCX reports from sensor data"""
    
    def __init__(self, template_path: str):
        """
        Initialize DOCX generator
        
        Args:
            template_path: Path to template DOCX file
        """
        self.template_path = template_path
        self.date_formatter = DateFormatter('en')
    
    def create_document_from_template(self, output_path: str) -> Document:
        """
        Create new document based on template
        
        Args:
            output_path: Path for output file
            
        Returns:
            Document object
        """
        # Copy template to output path
        shutil.copy2(self.template_path, output_path)
        
        # Open the copied document
        return Document(output_path)
    
    def set_cell_background_color(self, cell, color: str):
        """
        Set background color of table cell
        
        Args:
            cell: Table cell object
            color: Hex color code (e.g., 'FF0000' for red)
        """
        try:
            # Get cell properties
            tc_pr = cell._tc.get_or_add_tcPr()
            
            # Create shading element
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color)
            shd.set(qn('w:val'), 'clear')
            
            # Add to cell properties
            tc_pr.append(shd)
        except Exception as e:
            print(f"Warning: Could not set cell background color: {e}")
    
    def set_cell_vertical_alignment(self, cell, alignment=WD_ALIGN_VERTICAL.CENTER):
        """
        Set vertical alignment for a table cell
        
        Args:
            cell: Table cell object
            alignment: Vertical alignment (WD_ALIGN_VERTICAL.CENTER, TOP, BOTTOM)
        """
        try:
            # Get cell properties
            tc_pr = cell._tc.get_or_add_tcPr()
            
            # Remove existing vertical alignment if present
            existing_valign = tc_pr.find(qn('w:vAlign'))
            if existing_valign is not None:
                tc_pr.remove(existing_valign)
            
            # Create vertical alignment element
            valign = OxmlElement('w:vAlign')
            
            # Map alignment enum to string value
            alignment_map = {
                WD_ALIGN_VERTICAL.CENTER: 'center',
                WD_ALIGN_VERTICAL.TOP: 'top',
                WD_ALIGN_VERTICAL.BOTTOM: 'bottom'
            }
            
            valign.set(qn('w:val'), alignment_map.get(alignment, 'center'))
            
            # Add to cell properties
            tc_pr.append(valign)
        except Exception as e:
            print(f"Warning: Could not set cell vertical alignment: {e}")
    
    def format_temperature(self, temperature: float) -> str:
        """Format temperature according to %0.1f"""
        return f"{temperature:.1f}"
    
    def format_humidity(self, humidity: float) -> str:
        """Format humidity according to %0.0f"""
        return f"{humidity:.0f}"
    
    def get_compliance_color(self, is_compliant: bool) -> str:
        """
        Get background color for compliance status
        
        Args:
            is_compliant: Whether the value is compliant
            
        Returns:
            Hex color code
        """
        if is_compliant:
            return 'C6EFCE'  # Light green
        else:
            return 'FFC7CE'  # Light red
    
    def get_week_background_color(self, date, month: int, year: int) -> str:
        """
        Get alternating background color for weeks
        
        Args:
            date: Date object
            month: Month number
            year: Year
            
        Returns:
            Hex color code for week background
        """
        # Get week number within the month
        # Start from the first day of the month
        first_day = datetime.date(year, month, 1)
        first_week = first_day.isocalendar().week
        current_week = date.isocalendar().week
        
        # Calculate week offset from first week of month
        week_offset = current_week - first_week
        
        # Handle year boundary (e.g., January might have week 52/53 from previous year)
        if week_offset < 0:
            week_offset += 53  # Assume max 53 weeks per year
        
        # Alternate colors: even weeks white, odd weeks light gray
        if week_offset % 2 == 0:
            return 'FFFFFF'  # White
        else:
            return 'F5F5F5'  # Light gray
    
    def add_data_to_table(self, doc: Document, daily_data: List[Dict], 
                         month: int, year: int) -> bool:
        """
        Add sensor data to the document table
        
        Args:
            doc: Document object
            daily_data: List of daily processed data
            month: Month number
            year: Year
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Find the table in the document
            if not doc.tables:
                print("No tables found in document")
                return False
            
            table = doc.tables[0]  # Assume first table is our data table
            
            # Clear existing data rows (keep header)
            # Start from row 1 (index 1) to preserve header
            rows_to_remove = []
            for i in range(1, len(table.rows)):
                rows_to_remove.append(i)
            
            # Remove rows in reverse order to maintain indices
            for row_idx in reversed(rows_to_remove):
                table._element.remove(table.rows[row_idx]._element)
            
            # Add data rows
            for day_data in daily_data:
                date = day_data['date']
                morning = day_data['morning']
                evening = day_data['evening']
                
                # Format date string
                date_str = self.date_formatter.format_date(date, 'ddd dd/MM/yy')
                
                # Always add both morning (08:00) and evening (16:00) rows for each day
                # This ensures proper table structure and cell merging
                
                # Add morning row (08:00) - always first
                morning_row = table.add_row()
                
                if morning:
                    # Morning data available
                    morning_row.cells[0].text = self.format_humidity(morning['humidity'])
                    morning_row.cells[1].text = self.format_temperature(morning['temperature'])
                    
                    # Set compliance colors
                    humidity_color = self.get_compliance_color(morning['humidity_compliant'])
                    temp_color = self.get_compliance_color(morning['temp_compliant'])
                    self.set_cell_background_color(morning_row.cells[0], humidity_color)
                    self.set_cell_background_color(morning_row.cells[1], temp_color)
                else:
                    # No morning data
                    morning_row.cells[0].text = "-"
                    morning_row.cells[1].text = "-"
                    self.set_cell_background_color(morning_row.cells[0], 'F2F2F2')  # Light gray
                    self.set_cell_background_color(morning_row.cells[1], 'F2F2F2')  # Light gray
                
                # Set time and date for morning row
                morning_row.cells[2].text = "08:00"
                morning_row.cells[3].text = date_str
                
                # Get week background color for this date
                week_color = self.get_week_background_color(date, month, year)
                
                # Set week background color for time and date cells
                self.set_cell_background_color(morning_row.cells[2], week_color)  # Time cell
                self.set_cell_background_color(morning_row.cells[3], week_color)  # Date cell
                
                # Set alignment for morning row
                for cell in morning_row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add evening row (16:00) - always second
                evening_row = table.add_row()
                
                if evening:
                    # Evening data available
                    evening_row.cells[0].text = self.format_humidity(evening['humidity'])
                    evening_row.cells[1].text = self.format_temperature(evening['temperature'])
                    
                    # Set compliance colors
                    humidity_color = self.get_compliance_color(evening['humidity_compliant'])
                    temp_color = self.get_compliance_color(evening['temp_compliant'])
                    self.set_cell_background_color(evening_row.cells[0], humidity_color)
                    self.set_cell_background_color(evening_row.cells[1], temp_color)
                else:
                    # No evening data
                    evening_row.cells[0].text = "-"
                    evening_row.cells[1].text = "-"
                    self.set_cell_background_color(evening_row.cells[0], 'F2F2F2')  # Light gray
                    self.set_cell_background_color(evening_row.cells[1], 'F2F2F2')  # Light gray
                
                # Set time for evening row (date cell will be merged)
                evening_row.cells[2].text = "16:00"
                evening_row.cells[3].text = ""  # Empty for merged cell effect
                
                # Set week background color for time cell (date cell will be merged)
                self.set_cell_background_color(evening_row.cells[2], week_color)  # Time cell
                
                # Set alignment for evening row
                for cell in evening_row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Merge date cells for the day (morning and evening rows)
                try:
                    # Get the date cell from morning row and merge with evening row date cell
                    morning_date_cell = morning_row.cells[3]
                    evening_date_cell = evening_row.cells[3]
                    
                    # Merge cells vertically
                    morning_date_cell.merge(evening_date_cell)
                    
                    # Set vertical alignment for the merged date cell
                    self.set_cell_vertical_alignment(morning_date_cell, WD_ALIGN_VERTICAL.CENTER)
                    
                except Exception as e:
                    print(f"Warning: Could not merge date cells: {e}")
                    # If merging fails, at least put the date in both cells
                    evening_row.cells[3].text = date_str
                    # Still try to set vertical alignment for both cells
                    self.set_cell_vertical_alignment(morning_row.cells[3], WD_ALIGN_VERTICAL.CENTER)
                    self.set_cell_vertical_alignment(evening_row.cells[3], WD_ALIGN_VERTICAL.CENTER)
            
            return True
            
        except Exception as e:
            print(f"Error adding data to table: {e}")
            return False
    
    def update_document_title(self, doc: Document, month: int, year: int, csv_filename: str = None) -> bool:
        """
        Update document title with month/year and optionally CSV filename
        
        Args:
            doc: Document object
            month: Month number
            year: Year
            csv_filename: Optional CSV filename (without extension) to add to title
            
        Returns:
            True if successful
        """
        try:
            # Find and update title paragraph
            for paragraph in doc.paragraphs:
                if "Temperature and humidity monitoring" in paragraph.text:
                    # Add month/year to title
                    month_name = self.date_formatter.get_month_name(month)
                    
                    # Build title with optional CSV filename
                    if csv_filename:
                        new_title = f"Temperature and humidity monitoring - {csv_filename} ({month_name} {year})"
                    else:
                        new_title = f"Temperature and humidity monitoring ({month_name} {year})"
                    
                    paragraph.text = new_title
                    return True
            
            return False
            
        except Exception as e:
            print(f"Error updating document title: {e}")
            return False
    
    def generate_report(self, daily_data: List[Dict], month: int, year: int, 
                       output_path: str, csv_filename: str = None) -> bool:
        """
        Generate complete DOCX report
        
        Args:
            daily_data: List of daily processed data
            month: Month number
            year: Year
            output_path: Path for output file
            csv_filename: Optional CSV filename (without extension) to add to title
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create document from template
            doc = self.create_document_from_template(output_path)
            
            # Update title with CSV filename
            self.update_document_title(doc, month, year, csv_filename)
            
            # Add data to table
            if not self.add_data_to_table(doc, daily_data, month, year):
                print("Failed to add data to table")
                return False
            
            # Save document
            doc.save(output_path)
            print(f"Report generated successfully: {output_path}")
            return True
            
        except Exception as e:
            print(f"Error generating report: {e}")
            return False


if __name__ == "__main__":
    # Test the DOCX generator
    import datetime
    
    # Create test data
    test_daily_data = [
        {
            'date': datetime.date(2025, 8, 21),
            'morning': {
                'temperature': 24.2,
                'humidity': 55.0,
                'temp_compliant': True,
                'humidity_compliant': True,
                'fully_compliant': True,
                'readings_count': 3
            },
            'evening': {
                'temperature': 25.8,
                'humidity': 48.0,
                'temp_compliant': True,
                'humidity_compliant': True,
                'fully_compliant': True,
                'readings_count': 2
            }
        },
        {
            'date': datetime.date(2025, 8, 22),
            'morning': {
                'temperature': 27.1,  # Non-compliant (> 26)
                'humidity': 35.0,     # Non-compliant (< 40)
                'temp_compliant': False,
                'humidity_compliant': False,
                'fully_compliant': False,
                'readings_count': 2
            },
            'evening': None  # No evening data
        }
    ]
    
    # Test generator (requires template.docx to exist)
    if os.path.exists('template.docx'):
        generator = DOCXGenerator('template.docx')
        success = generator.generate_report(test_daily_data, 8, 2025, 'test_report.docx')
        print(f"Test generation successful: {success}")
    else:
        print("Template file not found for testing")